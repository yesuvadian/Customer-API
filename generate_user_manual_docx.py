#!/usr/bin/env python3
"""
Generate User Manual DOCX
Converts the user manual into a professional DOCX document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_heading(doc, text, level=1):
    """Add a heading with custom formatting"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_table_with_data(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Bold header text
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

def create_user_manual():
    """Create the user manual DOCX document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ========================================
    # TITLE PAGE
    # ========================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CogniWatt Customer Portal\n')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('User Manual')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()  # Spacing

    version_info = doc.add_paragraph()
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_info.add_run(f'Version 1.0\n').font.size = Pt(14)
    version_info.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n').font.size = Pt(12)
    version_info.add_run('Karnataka Power Transmission Corporation Limited').font.size = Pt(12)

    doc.add_page_break()

    # ========================================
    # TABLE OF CONTENTS
    # ========================================
    add_heading(doc, 'Table of Contents', 1)

    toc_items = [
        '1. System Overview',
        '2. Login Credentials',
        '3. User Roles & Permissions',
        '4. Getting Started',
        '5. Testing Request Workflow',
        '6. User Guides by Role',
        '7. Department Hierarchy',
        '8. Auto-Assignment System',
        '9. Common Workflows',
        '10. Troubleshooting',
        '11. Best Practices',
        '12. Support & Contact'
    ]

    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)

    doc.add_page_break()

    # ========================================
    # 1. SYSTEM OVERVIEW
    # ========================================
    add_heading(doc, '1. System Overview', 1)

    doc.add_paragraph(
        'CogniWatt Customer Portal is an equipment testing management system for KPTCL '
        '(Karnataka Power Transmission Corporation Limited). The system manages the complete '
        'lifecycle of testing requests from submission to approval with automated tester assignment.'
    )

    doc.add_heading('Key Features', 2)
    features = [
        'Multi-tenant organization support',
        'Hierarchical department structure (6 levels)',
        'Role-based access control',
        'Automated tester assignment with workload balancing',
        'Complete workflow engine with 9 states',
        'Testing request tracking and approval',
        'Department-scoped permissions'
    ]
    for feature in features:
        doc.add_paragraph(f'✓ {feature}', style='List Bullet')

    doc.add_page_break()

    # ========================================
    # 2. LOGIN CREDENTIALS
    # ========================================
    add_heading(doc, '2. Login Credentials', 1)

    doc.add_paragraph(
        'All users have the default password: admin123\n\n'
        '⚠️ Security Note: Please change your password after first login.',
        style='Intense Quote'
    )

    doc.add_heading('Sample User Accounts', 2)

    headers = ['Email', 'Role', 'Department', 'Password']
    rows = [
        ['orgadmin@kptcl.com', 'Organization Admin', 'Organization Level', 'admin123'],
        ['depthead@kptcl.com', 'Department Head', 'RT North Division', 'admin123'],
        ['tester1@kptcl.com', 'Tester', 'Yelahanka Section', 'admin123'],
        ['tester2@kptcl.com', 'Tester', 'RT North SD1 Yelahanka', 'admin123'],
        ['engineer@kptcl.com', 'Engineer', '220kV Yelahanka Substation', 'admin123']
    ]

    add_table_with_data(doc, headers, rows)

    doc.add_page_break()

    # ========================================
    # 3. USER ROLES & PERMISSIONS
    # ========================================
    add_heading(doc, '3. User Roles & Permissions', 1)

    # Organization Admin
    doc.add_heading('3.1 Organization Admin', 2)
    doc.add_paragraph('Email: orgadmin@kptcl.com', style='List Bullet')
    doc.add_paragraph('Scope: Full organization access', style='List Bullet')

    doc.add_paragraph('\nPermissions:')
    admin_perms = [
        'View all departments and users',
        'Create/edit departments at all levels',
        'Manage roles and permissions',
        'View all testing requests across the organization',
        'Assign/reassign testers',
        'Configure workflow and permission matrix',
        'Access analytics and reports'
    ]
    for perm in admin_perms:
        doc.add_paragraph(f'✓ {perm}', style='List Bullet 2')

    # Department Head
    doc.add_heading('3.2 Department Head', 2)
    doc.add_paragraph('Email: depthead@kptcl.com', style='List Bullet')
    doc.add_paragraph('Department: RT North Division', style='List Bullet')
    doc.add_paragraph('Scope: Department tree access', style='List Bullet')

    doc.add_paragraph('\nPermissions:')
    dept_perms = [
        'View all requests in their department hierarchy',
        'Approve/reject test results',
        'View workload statistics for testers',
        'Assign/reassign testers within department scope',
        'View department dashboard',
        'Manage users within department'
    ]
    for perm in dept_perms:
        doc.add_paragraph(f'✓ {perm}', style='List Bullet 2')

    # Tester
    doc.add_heading('3.3 Tester', 2)
    doc.add_paragraph('Emails: tester1@kptcl.com, tester2@kptcl.com', style='List Bullet')
    doc.add_paragraph('Scope: Department-specific access', style='List Bullet')

    doc.add_paragraph('\nPermissions:')
    tester_perms = [
        'View assigned testing requests',
        'Accept/reject testing assignments',
        'Start testing',
        'Submit test results',
        'Upload test reports and images',
        'View testing history',
        'Update testing status'
    ]
    for perm in tester_perms:
        doc.add_paragraph(f'✓ {perm}', style='List Bullet 2')

    # Engineer
    doc.add_heading('3.4 Engineer', 2)
    doc.add_paragraph('Email: engineer@kptcl.com', style='List Bullet')
    doc.add_paragraph('Department: 220kV Yelahanka Substation', style='List Bullet')
    doc.add_paragraph('Scope: Department-specific access', style='List Bullet')

    doc.add_paragraph('\nPermissions:')
    engineer_perms = [
        'Create testing requests',
        'View own testing requests',
        'View test results and reports',
        'Cancel draft requests',
        'Track request status',
        'View testing history'
    ]
    for perm in engineer_perms:
        doc.add_paragraph(f'✓ {perm}', style='List Bullet 2')

    doc.add_page_break()

    # ========================================
    # 4. GETTING STARTED
    # ========================================
    add_heading(doc, '4. Getting Started', 1)

    doc.add_heading('First Login', 2)

    doc.add_paragraph('Step 1: Access the Portal')
    doc.add_paragraph('URL: http://localhost:8000 (or your deployment URL)', style='List Bullet')

    doc.add_paragraph('\nStep 2: Login Steps')
    login_steps = [
        'Enter your email address',
        'Enter password: admin123',
        'Click "Login"'
    ]
    for i, step in enumerate(login_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('\nStep 3: Change Password (Recommended)')
    password_steps = [
        'Go to Profile → Security Settings',
        'Click "Change Password"',
        'Enter current password and new password',
        'Save changes'
    ]
    for i, step in enumerate(password_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_page_break()

    # ========================================
    # 5. TESTING REQUEST WORKFLOW
    # ========================================
    add_heading(doc, '5. Testing Request Workflow', 1)

    doc.add_heading('Workflow States', 2)

    doc.add_paragraph(
        'Draft → Submitted → Assigned → Accepted → In Progress → '
        'Test Submitted → Approved'
    )

    doc.add_heading('State Descriptions', 2)

    headers = ['State', 'Description', 'Who Can See']
    rows = [
        ['Draft', 'Request being prepared', 'Requester only'],
        ['Submitted', 'Awaiting tester assignment', 'Requester, Dept Head'],
        ['Assigned', 'Tester assigned', 'Tester, Requester, Dept Head'],
        ['Accepted', 'Tester accepted assignment', 'Tester, Requester, Dept Head'],
        ['In Progress', 'Testing underway', 'Tester, Requester, Dept Head'],
        ['Test Submitted', 'Results awaiting approval', 'Tester, Dept Head, Admin'],
        ['Approved', 'Final approval (completed)', 'All'],
        ['Rejected', 'Rejected at any stage', 'All'],
        ['Cancelled', 'Cancelled by requester', 'All']
    ]

    add_table_with_data(doc, headers, rows)

    doc.add_page_break()

    # ========================================
    # 6. USER GUIDES BY ROLE
    # ========================================
    add_heading(doc, '6. User Guides by Role', 1)

    # Engineer Guide
    doc.add_heading('6.1 Engineer Guide: Creating a Testing Request', 2)

    doc.add_paragraph('Login: engineer@kptcl.com / admin123', style='Intense Quote')

    doc.add_paragraph('\nStep 1: Navigate to Create Request')
    doc.add_paragraph('1. Login to the portal', style='List Number')
    doc.add_paragraph('2. Go to Testing Requests → Create New Request', style='List Number')

    doc.add_paragraph('\nStep 2: Fill Request Details')
    doc.add_paragraph('Equipment Details:', style='Heading 3')
    equipment_fields = [
        'Equipment Type: Select from dropdown',
        'Test Type: Select test to perform',
        'Transformer Type: Power/Distribution',
        'Transformer Rating: e.g., 100 MVA, 11/0.433 kV',
        'Manufacturer: Equipment manufacturer name',
        'Serial Number: Manufacturer serial number'
    ]
    for field in equipment_fields:
        doc.add_paragraph(field, style='List Bullet 2')

    doc.add_paragraph('\nRequest Information:', style='Heading 3')
    request_fields = [
        'Title: Brief description',
        'Description: Detailed testing requirements',
        'Priority: Normal / High / Urgent',
        'Requested Date: When testing should be done',
        'Due Date: Deadline for completion'
    ]
    for field in request_fields:
        doc.add_paragraph(field, style='List Bullet 2')

    doc.add_paragraph('\nStep 3: Submit Request')
    doc.add_paragraph(
        'Click "Submit Request" button. No need to select a tester - '
        'the system will auto-assign based on workload and availability!',
        style='Intense Quote'
    )

    # Tester Guide
    doc.add_page_break()
    doc.add_heading('6.2 Tester Guide: Processing Testing Requests', 2)

    doc.add_paragraph('Login: tester1@kptcl.com or tester2@kptcl.com / admin123', style='Intense Quote')

    doc.add_paragraph('\nStep 1: View Assigned Requests')
    view_steps = [
        'Login to portal',
        'Go to Dashboard → Assigned to Me',
        'View list of pending assignments'
    ]
    for i, step in enumerate(view_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('\nStep 2: Accept Assignment')
    accept_steps = [
        'Click on the assigned request',
        'Review equipment details and requirements',
        'Click "Accept Assignment"',
        'Or click "Reject Assignment" with reason if unable to test'
    ]
    for i, step in enumerate(accept_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('\nStep 3: Start Testing')
    doc.add_paragraph('Once accepted, click "Start Testing"', style='List Bullet')
    doc.add_paragraph('Status changes to "In Progress"', style='List Bullet')
    doc.add_paragraph('Conduct the equipment tests as per specifications', style='List Bullet')

    doc.add_paragraph('\nStep 4: Submit Test Results')
    submit_steps = [
        'Click "Submit Test Results"',
        'Fill in test date, location, equipment used',
        'Enter test parameters and observations',
        'Upload test report PDF and images',
        'Select recommendation: Pass / Fail / Conditional / Retest',
        'Add notes/comments',
        'Click "Submit Results"'
    ]
    for i, step in enumerate(submit_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    # Department Head Guide
    doc.add_page_break()
    doc.add_heading('6.3 Department Head Guide: Approval Process', 2)

    doc.add_paragraph('Login: depthead@kptcl.com / admin123', style='Intense Quote')

    doc.add_paragraph('\nStep 1: View Pending Approvals')
    dept_steps = [
        'Login to portal',
        'Go to Dashboard → Pending Approvals',
        'View all test results awaiting approval'
    ]
    for i, step in enumerate(dept_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('\nStep 2: Review Test Results')
    doc.add_paragraph('Click on the request and review:', style='List Bullet')
    review_items = [
        'Equipment details',
        'Test parameters',
        'Test results',
        'Uploaded reports and images',
        'Tester recommendations'
    ]
    for item in review_items:
        doc.add_paragraph(f'✓ {item}', style='List Bullet 2')

    doc.add_paragraph('\nStep 3: Approve or Reject')
    doc.add_paragraph('Option A: Approve')
    approve_steps = [
        'Click "Approve"',
        'Add approval comments (optional)',
        'Click "Confirm"',
        'Status changes to "Approved"'
    ]
    for i, step in enumerate(approve_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_paragraph('\nOption B: Reject')
    reject_steps = [
        'Click "Reject"',
        'Enter rejection reason (mandatory)',
        'Specify what needs to be corrected',
        'Click "Confirm"',
        'Status changes to "Rejected"',
        'Tester is notified'
    ]
    for i, step in enumerate(reject_steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_page_break()

    # ========================================
    # 7. DEPARTMENT HIERARCHY
    # ========================================
    add_heading(doc, '7. Department Hierarchy', 1)

    doc.add_heading('KPTCL Organization Structure', 2)

    hierarchy = [
        'Karnataka Power Transmission Corporation Limited (KPTCL)',
        '  └── Bangalore Zone',
        '      └── Bangalore Transmission Circle',
        '          └── RT North Division',
        '              └── RT North SD1 Yelahanka (Subdivision)',
        '                  └── Yelahanka Section',
        '                      └── 220kV Yelahanka Substation'
    ]

    for level in hierarchy:
        p = doc.add_paragraph(level)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_heading('Department Types', 2)

    headers = ['Level', 'Type', 'Code', 'Example']
    rows = [
        ['1', 'Zone', 'ZONE', 'Bangalore Zone'],
        ['2', 'Circle', 'CIRCLE', 'Bangalore Transmission Circle'],
        ['3', 'Division', 'DIVISION', 'RT North Division'],
        ['4', 'Subdivision', 'SUBDIVISION', 'RT North SD1 Yelahanka'],
        ['5', 'Section', 'SECTION', 'Yelahanka Section'],
        ['6', 'Substation', 'SUBSTATION', '220kV Yelahanka Substation']
    ]

    add_table_with_data(doc, headers, rows)

    doc.add_page_break()

    # ========================================
    # 8. AUTO-ASSIGNMENT SYSTEM
    # ========================================
    add_heading(doc, '8. Auto-Assignment System', 1)

    doc.add_heading('How Tester Auto-Assignment Works', 2)

    doc.add_paragraph('When an engineer submits a testing request:')

    doc.add_paragraph('\n1. System analyzes eligible testers:')
    eligible_criteria = [
        'Same organization (KPTCL)',
        'Has "Tester" role',
        'Department hierarchy matching',
        'Active status'
    ]
    for criterion in eligible_criteria:
        doc.add_paragraph(f'• {criterion}', style='List Bullet 2')

    doc.add_paragraph('\n2. Workload balancing strategies:')

    strategies = [
        ('Least Loaded (Default)', 'Assigns to tester with fewest active requests'),
        ('Round Robin', 'Rotates assignments among testers'),
        ('Priority-Based', 'Assigns based on tester skill/experience'),
        ('Random', 'Random selection from eligible testers')
    ]

    for strategy, desc in strategies:
        p = doc.add_paragraph()
        p.add_run(f'{strategy}: ').bold = True
        p.add_run(desc)
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_paragraph('\n3. Availability check:')
    availability = [
        'Max concurrent assignments: 5 (configurable)',
        'Checks tester is not on leave',
        'Verifies active status'
    ]
    for item in availability:
        doc.add_paragraph(f'• {item}', style='List Bullet 2')

    doc.add_paragraph('\n4. Automatic assignment:')
    assignment_flow = [
        'Tester auto-assigned to request',
        'State changes: Submitted → Assigned',
        'Tester receives notification',
        'Requester sees assigned tester'
    ]
    for item in assignment_flow:
        doc.add_paragraph(f'• {item}', style='List Bullet 2')

    doc.add_page_break()

    # ========================================
    # 9. COMMON WORKFLOWS
    # ========================================
    add_heading(doc, '9. Common Workflows', 1)

    doc.add_heading('Workflow 1: Successful Testing Request', 2)

    workflow1 = [
        'Engineer creates request',
        'Engineer submits request (Draft → Submitted)',
        'System auto-assigns tester (Submitted → Assigned)',
        'Tester accepts assignment (Assigned → Accepted)',
        'Tester starts testing (Accepted → In Progress)',
        'Tester submits results (In Progress → Test Submitted)',
        'Dept Head approves (Test Submitted → Approved)',
        '✓ Request completed'
    ]

    for i, step in enumerate(workflow1, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        if '✓' in step:
            p.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            p.runs[0].bold = True

    doc.add_heading('Workflow 2: Request Rejected by Tester', 2)

    workflow2 = [
        'System assigns tester (Submitted → Assigned)',
        'Tester rejects with reason (Assigned → Rejected)',
        'System auto-reassigns to next available tester',
        'New tester accepts...'
    ]

    for i, step in enumerate(workflow2, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ========================================
    # 10. TROUBLESHOOTING
    # ========================================
    add_heading(doc, '10. Troubleshooting', 1)

    issues = [
        ('Cannot Login', [
            'Check email is correct (including @kptcl.com domain)',
            'Verify password is admin123 (case-sensitive)',
            'Check CAPS LOCK is off',
            'Clear browser cache and cookies',
            'Contact admin to verify account is active'
        ]),
        ('Cannot See Testing Requests', [
            'Check you have correct role permissions',
            'Verify department assignment',
            'Check filter settings (status, date range)',
            'Engineers: Go to "My Requests" tab',
            'Testers: Go to "Assigned to Me" tab'
        ]),
        ('Auto-Assignment Not Working', [
            'Verify testers exist in same organization',
            'Check testers have "Tester" role assigned',
            'Ensure testers are in matching department hierarchy',
            'Verify testers are active (not on leave)',
            'Check tester workload < 5 active requests',
            'Contact admin for manual assignment'
        ])
    ]

    for issue, solutions in issues:
        doc.add_heading(f'Issue: {issue}', 2)
        doc.add_paragraph('Solutions:')
        for solution in solutions:
            doc.add_paragraph(f'✓ {solution}', style='List Bullet')
        doc.add_paragraph()

    doc.add_page_break()

    # ========================================
    # 11. BEST PRACTICES
    # ========================================
    add_heading(doc, '11. Best Practices', 1)

    doc.add_heading('For Engineers (Requesters)', 2)

    doc.add_paragraph('DO:', style='Heading 3')
    do_items = [
        'Provide complete equipment details',
        'Set realistic due dates',
        'Add clear testing requirements in description',
        'Attach equipment specifications if available',
        'Track request progress regularly'
    ]
    for item in do_items:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')

    doc.add_paragraph('\nDON\'T:', style='Heading 3')
    dont_items = [
        'Submit incomplete requests',
        'Set impossible deadlines',
        'Create duplicate requests',
        'Cancel requests unnecessarily'
    ]
    for item in dont_items:
        doc.add_paragraph(f'✗ {item}', style='List Bullet')

    doc.add_heading('For Testers', 2)

    doc.add_paragraph('DO:', style='Heading 3')
    tester_do = [
        'Accept assignments promptly',
        'Start testing on scheduled date',
        'Document all test parameters',
        'Upload clear test reports',
        'Add detailed observations',
        'Submit results with recommendations'
    ]
    for item in tester_do:
        doc.add_paragraph(f'✓ {item}', style='List Bullet')

    doc.add_paragraph('\nDON\'T:', style='Heading 3')
    tester_dont = [
        'Accept more assignments than you can handle',
        'Delay testing without communication',
        'Submit incomplete test results',
        'Skip uploading test documentation'
    ]
    for item in tester_dont:
        doc.add_paragraph(f'✗ {item}', style='List Bullet')

    doc.add_page_break()

    # ========================================
    # 12. SUPPORT & CONTACT
    # ========================================
    add_heading(doc, '12. Support & Contact', 1)

    doc.add_heading('Technical Support', 2)
    support_info = [
        ('Email', 'support@cogniwatt.com'),
        ('Phone', '+91-XXXX-XXXXX'),
        ('Hours', 'Monday - Friday, 9:00 AM - 6:00 PM IST')
    ]
    for label, value in support_info:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)

    doc.add_heading('System Administrator', 2)
    admin_info = [
        ('Name', 'Organization Admin'),
        ('Email', 'orgadmin@kptcl.com'),
        ('For', 'User account issues, permissions, system configuration')
    ]
    for label, value in admin_info:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)

    doc.add_heading('Training & Documentation', 2)
    doc_links = [
        'User Manual: This document',
        'Video Tutorials: Available on training portal',
        'FAQ: Available on support portal',
        'Release Notes: Check for updates'
    ]
    for link in doc_links:
        doc.add_paragraph(link, style='List Bullet')

    doc.add_page_break()

    # ========================================
    # DOCUMENT INFORMATION
    # ========================================
    add_heading(doc, 'Document Information', 1)

    doc_info = [
        ('Version', '1.0'),
        ('Last Updated', datetime.now().strftime('%B %d, %Y')),
        ('Document Owner', 'CogniWatt Development Team'),
        ('Next Review Date', 'June 22, 2026'),
        ('Organization', 'Karnataka Power Transmission Corporation Limited (KPTCL)')
    ]

    for label, value in doc_info:
        p = doc.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(value)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('—' * 50)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    end_note = doc.add_paragraph('End of User Manual')
    end_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_note.runs[0].font.italic = True

    # Save document
    output_path = 'CogniWatt_User_Manual.docx'
    doc.save(output_path)
    print(f'[OK] User manual generated: {output_path}')
    return output_path

if __name__ == '__main__':
    import os
    try:
        output_file = create_user_manual()
        print(f'\n[SUCCESS] User manual created at: {output_file}')
        print(f'[INFO] File size: {os.path.getsize(output_file) / 1024:.2f} KB')
    except Exception as e:
        print(f'[ERROR] Error generating user manual: {e}')
        import traceback
        traceback.print_exc()
